from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from docx import Document
import io
import os
import shutil
import sqlite3
from datetime import datetime, timedelta, timezone
import firebase_admin
from firebase_admin import credentials, db as admin_db
import uuid
import ipaddress
import subprocess
import platform
import socket
import re
import requests
import urllib3
import hashlib
from flask_cors import CORS
from urllib.parse import urlparse
from dateutil import parser
import smtplib
from email.mime.text import MIMEText


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

app = Flask(__name__)
app.secret_key = 'HelloWorld'
CORS(app, origins=["chrome-extension://fclopfkjjjeadnfknjdcnfljinibpmbg", "http://localhost:5000"])


# Firebase Admin SDK initialization
cred = credentials.Certificate('goldenshield-01-firebase-adminsdk-fbsvc-cc0ff4e580.json')
firebase_admin.initialize_app(cred, {
    'databaseURL': "https://goldenshield-01-default-rtdb.asia-southeast1.firebasedatabase.app"
})


db = firebase_admin.db

def get_current_ip():
    """Get the current IP address of the wireless LAN adapter."""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        current_ip = s.getsockname()[0]
        s.close()
        return current_ip
    except Exception as e:
        print(f"Error getting current IP: {e}")
        return None


def get_allowed_network():
    """Get the allowed network based on the current IP address."""
    current_ip = get_current_ip()
    if current_ip:
        return ipaddress.ip_network(f"{current_ip.rsplit('.', 1)[0]}.0/24")
    return None


def is_allowed_ip(ip):
    """Allow access from any IP address."""
    # Allow localhost
    if ip in ["127.0.0.1", "::1"]:
        return True
    return True


def get_mac_address(ip_address):
    """Get the MAC address of a device from its IP on local network using the ARP table."""
    try:
        if platform.system().lower() == "windows":
            output = subprocess.check_output(f"arp -a {ip_address}", shell=True).decode()
            mac_match = re.search(r"([0-9a-f]{2}-){5}[0-9a-f]{2}", output.lower())
        else:
            output = subprocess.check_output(["arp", "-n", ip_address]).decode()
            mac_match = re.search(r"([0-9a-f]{2}:){5}[0-9a-f]{2}", output.lower())

        if mac_match:
            return mac_match.group(0)
        return "MAC Not Found"
    except Exception as e:
        return f"Error: {str(e)}"


def get_device_name_from_router(ip_address):
    router_url = "https://192.168.0.1"
    session = requests.Session()
    session.verify = False
    password = "Admin12345"
    username = "admin"

    try:
        login_url = f"{router_url}/cgi-bin/luci/api/xqsystem/login"
        login_payload = {
            "username": username,
            "password": password
        }

        headers = {
            "Content-Type": "application/json"
        }

        login_resp = session.post(login_url, json=login_payload, headers=headers, timeout=5)
        print("[DEBUG] Login response text:", login_resp.text)

        if not login_resp.ok or not login_resp.text:
            print("[!] Login request failed or returned no content.")
            return "Unknown Device"

        try:
            login_data = login_resp.json()
        except Exception as json_err:
            print(f"[!] Failed to parse login JSON: {json_err}")
            return "Unknown Device"

        stok = login_data.get("token") or login_data.get("stok")
        if not stok:
            print("[!] Login failed: No session token received.")
            return "Unknown Device"

        device_url = f"{router_url}/cgi-bin/luci/;stok={stok}/api/misystem/devicelist"
        device_resp = session.get(device_url, timeout=5)
        if not device_resp.ok or not device_resp.text:
            print("[!] Device list request failed or returned no content.")
            return "Unknown Device"

        try:
            device_data = device_resp.json()
        except Exception as json_err:
            print(f"[!] Failed to parse device list JSON: {json_err}")
            return "Unknown Device"

        for device in device_data.get("list", []):
            if device.get("ip") == ip_address:
                return device.get("name", "Unknown Device")

        return "Unknown Device"

    except Exception as e:
        print(f"[!] Error fetching device name: {e}")
        return "Unknown Device"


def tp_link_login(session, router_url, password):
    # Step 1: Get the login token (nonce)
    r = session.get(f"{router_url}/cgi-bin/luci/")
    nonce_match = re.search(r'"token"\s*:\s*"(\w+)"', r.text)
    if not nonce_match:
        print("[!] Failed to get login token")
        return None
    token = nonce_match.group(1)

    # Step 2: Hash password with token: md5(token + md5(password))
    md5_pass = hashlib.md5(password.encode()).hexdigest()
    hash_pass = hashlib.md5((token + md5_pass).encode()).hexdigest()

    # Step 3: Send login POST request
    login_payload = {
        "method": "do",
        "login": {
            "password": hash_pass
        }
    }

    login_url = f"{router_url}/cgi-bin/luci/;stok=/login?form=login"
    login_resp = session.post(login_url, json=login_payload, timeout=5)
    login_data = login_resp.json()
    stok = login_data.get('stok')

    if not stok:
        print("[!] Login failed: No session token received.")
        return None
    return stok


def fetch_tp_link_devices():
    router_url = "https://tplinkwifi.net/webpages/index.html?t=e6bfcc9c"
    password = "Admin12345"

    session = requests.Session()
    session.verify = False

    try:
        login_payload = {
            "method": "do",
            "login": {
                "password": password
            }
        }

        login_url = f"{router_url}/cgi-bin/luci/;stok=/login?form=login"
        login_resp = session.post(login_url, json=login_payload, timeout=5)

        if login_resp.status_code != 200:
            print("[!] Login failed: No session token received.")
            return []

        login_data = login_resp.json()
        stok = login_data.get('stok')

        if not stok:
            print("[!] Login failed: No session token received.")
            return []

        # Step 2: Fetch the device list
        device_url = f"{router_url}/cgi-bin/luci/;stok={stok}/api/misystem/devicelist"
        device_resp = session.get(device_url, timeout=5)

        if device_resp.status_code != 200:
            print("[!] Device list request failed or returned no content.")
            return []

        device_data = device_resp.json()

        connected_devices = []
        for device in device_data.get("list", []):
            name = device.get("name", "Unknown")
            ip = device.get("ip", "N/A")
            mac = device.get("mac", "N/A")
            connected_devices.append({
                'name': name,
                'ip': ip,
                'mac': mac
            })

        return connected_devices

    except Exception as e:
        print(f"[!] Error: {e}")
        return []


@app.route('/')
def index():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    return render_template('landing.html')


@app.route('/landing')
def landing():
    return render_template('landing.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        flash("Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", "error")
        return redirect(url_for('login'))

    now_str = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
    admin_users = admin_db.reference('admin_users').get() or {}
    senior_users = admin_db.reference('senior_users').get() or {}

    # Initialize session variables for tracking login attempts
    if 'login_attempts' not in session:
        session['login_attempts'] = 0
        session['lock_time'] = None

    # Check if the account is locked
    if session['lock_time'] and datetime.now(timezone.utc) < session['lock_time']:
        remaining_time = (session['lock_time'] - datetime.now(timezone.utc)).total_seconds()
        flash(f"Your account is locked. Try again in {int(remaining_time)} seconds.", "error")
        return render_template('login.html', remaining_time=int(remaining_time))

    # Handle POST
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')

        # Admin login check
        for uid, user in admin_users.items():
            if user.get('email') == email and user.get('password') == password:
                session['user'] = {
                    'fullname': user.get('fullname'),
                    'role': user.get('role'),
                    'email': user.get('email'),
                    'uid': uid
                }
                admin_db.reference(f'admin_users/{uid}').update({
                    "status": "online",
                    "last_active": now_str
                })
                session['login_attempts'] = 0  # Reset attempts on successful login
                return redirect(url_for('admin_dashboard'))

        # Senior login check
        for uid, user in senior_users.items():
            if user.get('email') == email and user.get('password') == password:
                session['user'] = {
                    'fullname': user.get('fullname'),
                    'role': user.get('role'),
                    'email': user.get('email'),
                    'uid': uid
                }

                # Initialize connection duration to zero if not already set
                connection_duration = admin_db.reference(f'senior_users/{uid}/connection_duration').get()
                if connection_duration is None:
                    admin_db.reference(f'senior_users/{uid}').update({
                        "connection_duration": {
                            "hours": 0,
                            "minutes": 0,
                            "seconds": 0
                        }
                    })

                admin_db.reference(f'senior_users/{uid}').update({
                    "status": "online",
                    "last_active": now_str
                })
                # Start tracking connection duration
                session['login_time'] = datetime.now(timezone.utc)
                session['login_attempts'] = 0  # Reset attempts on successful login
                return redirect(url_for('user_dashboard'))

        # Increment login attempts
        session['login_attempts'] += 1

        # Log suspicious activity to Firebase
        if session['login_attempts'] == 5:
            log_suspicious_activity(email, now_str)
        elif session['login_attempts'] >= 10:
            log_suspicious_activity(email, now_str)

        if session['login_attempts'] >= 10:
            session['lock_time'] = datetime.now(timezone.utc) + timedelta(seconds=30)
            flash("Your account is locked for 30 seconds due to too many failed attempts.", "error")
            return render_template('login.html', remaining_time=30)

        if session['login_attempts'] >= 5:
            flash("Are you trying to access another account?", "error")
            return redirect(url_for('login'))

        flash("Invalid email or password.", "error")

    return render_template('login.html')



@app.route('/logout')
def logout():
    if 'user' in session:
        user = session['user']
        uid = user.get('uid')
        role = user.get('role')
        now_str = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')

        # Update status to offline in Firebase
        if role == 'admin':
            admin_db.reference(f'admin_users/{uid}').update({
                "status": "offline",
                "last_active": now_str
            })
        elif role == 'senior':
            admin_db.reference(f'senior_users/{uid}').update({
                "status": "offline",
                "last_active": now_str
            })
            # Calculate connection duration
            if 'login_time' in session:
                login_time = session['login_time']
                connection_duration = datetime.now(timezone.utc) - login_time
                total_duration = admin_db.reference(f'senior_users/{uid}/connection_duration').get() or {}

                # Update the total connection duration
                total_duration_seconds = (total_duration.get('hours', 0) * 3600 +
                                          total_duration.get('minutes', 0) * 60 +
                                          total_duration.get('seconds', 0))
                total_duration_seconds += connection_duration.total_seconds()

                # Round total seconds to the nearest integer
                total_duration_seconds = round(total_duration_seconds)

                # Convert total seconds back to hours, minutes, seconds
                hours, remainder = divmod(total_duration_seconds, 3600)
                minutes, seconds = divmod(remainder, 60)

                admin_db.reference(f'senior_users/{uid}').update({
                    "connection_duration": {
                        "hours": hours,
                        "minutes": minutes,
                        "seconds": seconds
                    },
                    "status": "offline",
                    "last_active": now_str
                })

        session.pop('user', None)
        session.pop('login_time', None)  # Clear login time
        flash("Logged out successfully.", "success")

    return redirect(url_for('login'))


def load_blocked_adult_sites(file_path):
    blocked_sites = set()  # Use a set for faster lookups
    try:
        with open(file_path, 'r') as file:
            for line in file:
                url = line.strip().lower()  # Get the URL and convert to lowercase
                if url:  # Ensure the line is not empty
                    blocked_sites.add(url)
    except Exception as e:
        print(f"[ERROR] Failed to load blocked adult sites: {e}")
    return blocked_sites


# Load adult content sites
ADULT_CONTENT_FILE = 'adults_content/pornsite-list.txt'
BLOCKED_ADULT_SITES = load_blocked_adult_sites(ADULT_CONTENT_FILE)


# Define the list of blocked domains for gambling
BLOCKED_DOMAINS_GAMBLING = [
    "bingoplus.ph",
    "bingoplus.com",
    "bingoplus.net",
    "lucky.bingoplus.com",
    "fun.bingoplus.com",
    "play.bingoplus.com",
    "bingoplus.com.ph",
    "bingoplus.online",
    "bingoplusios1.com",
    "bingoplusios2.com",
    "bingoplusandroid1.com",
    "bingoplusandroid2.com",
    "fishplus.ph",
    "minesplus.ph",
    "blingwin.com",
    "happybingo.ph",
    "crazywin.ph",
    "747ph.live",
    "747.live",
    "gameph.com",
    "nustarmax.com",
    "patokbet.com",
    "pagcor.ph",
    "happyplay.ph",
    "gamefun.ph",
    "wwwgamefun.ph",
    "bp-poker.com",
    "bppoker.ph",
    "bppoker.net",
    "bppoker.live",
    "bppoker.app",
    "okbet.com",
    "okgames.com",
    "okgames.net",
    "ok-games.ph",
    "okapp.chat",
    "okbet.game",
    "okbet.help",
    "okbet.link",
    "okbet.net",
    "okbet.one",
    "okfun.ph",
    "okplay.ph",
    "okgames.ph",
    "hawkgaming.com",
    "s5.com",
    "ggpoker.ph",
    "filbet.com",
    "ps88.com",
    "pisogame.com",
    "pesogame.com",
    "buenas.ph",
    "inplay.ph",
    "igo.ph",
    "sportsplus.ph",
    "sportsplus.com.ph",
    "jadesportsbet.com",
    "sportsmaxx.ph",
    "fastwin.ph",
    "bethub.ph",
    "bethub.com.ph",
    "gamexsports.com",
    "gamexsports.ph",
    "gamexports.com.ph",
    "legendlink.com",
    "fbmemotion.ph",
    "fairplay.ph",
    "fairplay.com.ph",
    "ArionPlay.com",
    "bigwin29.com",
    "msw.ph",
    "egamescasino.ph",
    "playtime.ph",
    "playtime.com.ph",
    "spintime.ph",
    "cardtime.ph",
    "sulobet.ph",
    "LakiWin.com",
    "LakiPlay.com",
    "lets.playdailyfantasy.com",
    "deskgame.com",
    "deskgame.vip",
    "deskgame.org",
    "deskgame.co",
    "deskgame.club",
    "deskgame.me",
    "winzir.ph",
    "winzir.com",
    "winzir.net",
    "winzir.com.ph",
    "sg8.casino",
    "sg8.zone",
    "sg8.bet",
    "festival.sg8.casino",
    "bsports.ph",
    "arenaplus.ph",
    "arenaplus.net",
    "arenaplusvip.ph",
    "arenaplusvip.net",
    "arenaplusvip.com",
    "arena-plus.online",
    "arenaplusios1.com",
    "arenaplusios2.com",
    "arenaplusandroid1.com",
    "arenaplusandroid2.com",
    "arenaplus.asia",
    "arenaplus.info",
    "arenaplus.life",
    "arenaplus.org",
    "arenaplus.ph",
    "arenaplus.pro",
    "arenaplus.site",
    "arenaplus.today",
    "arenaplus.world",
    "arenapro.xyz",
    "arenaplus.fun",
    "gamezone88.com",
    "gamezone.ph",
    "gamezonebet.com",  "peryagame.com",  "peryagame.net",  "peryagame.ph","colorgameplus.com",
    "tripledg.com","tripledg1.com","tripledg2.com","bet88.ph","goplayasia.com","king.ph","lucky.ph", "queen.ph",
    "bosscat.ph",  "hqhole.com", "hdsexdino.com", "pornbigvideo.com", "sweetshow.com", "mylust.com", "sleazyneasy.com", "sexpulse.tv",
    "sexmole.com", "spankbang.com", "freeadultmedia.com", "bangbrosteenporn.com", "porn8.com", "collectionofbestporn.com",
    "hqporner.com", "freeviewmovies.com", "youporn.com", "pornhub.com", "adultfriendfinder.com", "adultfriendfinders.com",
    "adultfriendfinderz.com", "adultfriendsearch.com", "adultbanners.co.uk", "adultcash.com", "adultdatingtraffic.com",
    "adultmoneymakers.com", "adultpopunders.com", "adultrevenueservice.com", "adulttrafficads.com", "adultvalleycash.com",
    "adultwebmastersonline.com", "adultlinkexchange.com", "adultlinksco.com", "adultmoda.com", "adultadworld.com"
]


BLOCKED_URLS = []


def is_blocked_domain(domain):
    return domain in BLOCKED_DOMAINS_GAMBLING


def is_blocked_url(url):
    for blocked_url in BLOCKED_URLS:
        if blocked_url in url:
            return True
    return False


@app.before_request
def block_access():
    requested_url = request.url.lower()

    # Extract the domain from the requested URL
    parsed_url = urlparse(requested_url)
    domain = parsed_url.netloc

    # Fetch filtering preferences from Firebase
    filtering_preferences_ref = admin_db.reference('filtering_preferences')
    filtering_preferences = filtering_preferences_ref.get() or {}

    block_gambling = filtering_preferences.get('block_gambling', False)
    block_adult = filtering_preferences.get('block_adult', False)

    # Check if the requested domain is in the blocked adult sites list
    if block_adult and domain in BLOCKED_ADULT_SITES:
        return "Access Denied: Adult content sites are blocked.", 403

    if is_blocked_domain(domain) or is_blocked_url(requested_url):
        return "Access Denied: This site is blocked.", 403

    # Check if the requested domain is in the blocked gambling sites list
    if block_gambling and domain in BLOCKED_DOMAINS_GAMBLING:
        return "Access Denied: Gambling sites are blocked.", 403

    # Additional logic for blocked sites from Firebase (if needed)
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}
    blocked_list = [entry['website_url'].lower() for entry in blocked_sites.values()]

    # Check if the requested URL is in the blocked list
    for blocked_url in blocked_list:
        if blocked_url in requested_url:
            return "Access Denied: This site is blocked.", 403


@app.route('/check_blocked', methods=['POST'])
def check_blocked():
    data = request.get_json()
    website = data.get('website')

    if not website:
        return jsonify({"error": "Website URL is required."}), 400

    # Fetch blocked sites from Firebase
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}

    # Create a list of blocked site URLs
    blocked_list = [entry['website_url'].lower() for entry in blocked_sites.values()]

    # Check if the requested website is in the blocked list
    if website.lower() in blocked_list:
        # Log the attempt to access a blocked site
        user_session = session.get('user')
        if user_session:
            user_name = user_session.get('fullname', 'Unknown User')
            now_str = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')

            # Insert into Firebase notification_alerts
            notification_ref = admin_db.reference('notification_alerts')
            try:
                notification_ref.push({
                    'name': user_name,
                    'date': now_str.split(' ')[0],
                    'time': now_str.split(' ')[1],
                    'type': 'Attempting Blocked Sites',
                    'website': website,
                    'type_risk': 'High Risk'
                })
                print(f"[DEBUG] Logged blocked site access: {website} by {user_name}")
            except Exception as e:
                print(f"[ERROR] Failed to log blocked site access: {e}")

        return jsonify({"message": "Access denied. This site is blocked."}), 403

    return jsonify({"message": "Access granted."}), 200


def log_suspicious_activity(email, timestamp):
    """Log suspicious login attempts to Firebase."""
    notification_ref = admin_db.reference('notification_alerts')
    notification_ref.push({
        'name': email,  # Use the email as the name for the alert
        'date': timestamp.split(' ')[0],  # Extract date from timestamp
        'time': timestamp.split(' ')[1],  # Extract time from timestamp
        'type': 'Suspicious Activity',
        'website': 'N/A',
        'type_risk': 'Medium'
    })


@app.route('/create_account', methods=['GET'])
def create_account():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    return render_template('create_account.html')


@app.route('/register', methods=['POST'])
def register():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403

    mac_address = get_mac_address(client_ip)
    device_name = get_device_name_from_router(client_ip)

    fullname = request.form.get('fullname')
    email = request.form.get('email')
    password = request.form.get('password')
    role = request.form.get('role')
    status = "Offline"
    filter_level = "Easy"
    last_active = "Never"

    if role not in ['admin', 'senior']:
        flash("Invalid role selected.", "error")
        return redirect(url_for('create_account'))

    admin_users = admin_db.reference('admin_users').get() or {}
    senior_users = admin_db.reference('senior_users').get() or {}

    for user_dict in admin_users.values():
        if user_dict.get('email') == email:
            flash("Email already registered.", "error")
            return redirect(url_for('create_account'))

    for user_dict in senior_users.values():
        if user_dict.get('email') == email:
            flash("Email already registered.", "error")
            return redirect(url_for('create_account'))

    user_id = str(uuid.uuid4())

    user_data = {
        "fullname": fullname,
        "email": email,
        "password": password,
        "role": role,
        "status": status,
        "filter_level": filter_level,
        "last_active": last_active,
        "device": device_name,
        "ip_address": client_ip,
        "mac_address": mac_address
    }

    path = "admin_users" if role == 'admin' else "senior_users"
    admin_db.reference(f'{path}/{user_id}').set(user_data)

    flash("Account created successfully!", "success")
    return redirect(url_for('login'))


@app.route('/search', methods=['POST'])
def search():
    # Get the search query from the user
    data = request.get_json()
    search_query = data.get('query', '').strip().lower()

    # Fetch blocked sites from Firebase
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}

    # Create a list of blocked site URLs
    blocked_list = [entry['website_url'].lower() for entry in blocked_sites.values()]

    # Check if the search query is in the blocked list
    if search_query in blocked_list:
        return jsonify({"success": False, "message": "This site is blocked."}), 403

    # Proceed with the search logic (if not blocked)
    # For example, you can return a success message or perform a search operation
    return jsonify({"success": True, "message": "Search successful."})


@app.before_request
def block_access():
    # Get the requested URL
    requested_url = request.url.lower()

    # Fetch blocked sites from Firebase
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}

    # Create a list of blocked site URLs
    blocked_list = [entry['website_url'].lower() for entry in blocked_sites.values()]

    # Check if the requested URL is in the blocked list
    for blocked_url in blocked_list:
        if blocked_url in requested_url:
            return "Access Denied: This site is blocked.", 403


# ========== Admin Routes ==========

def check_admin_access():
    if 'user' not in session or session['user'].get('role') != 'admin':
        flash("Unauthorized access.", "error")
        return False
    return True


@app.route('/admin_dashboard')
def admin_dashboard():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session:
        return redirect(url_for('login'))
    user_role = session['user'].get('role')
    if user_role not in ['admin', 'Guardian']:
        flash("Unauthorized access.", "error")
        return redirect(url_for('login'))

    # Fetch senior users count
    senior_users_dict = admin_db.reference('senior_users').get() or {}
    senior_user_count = len(senior_users_dict)

    # Fetch blocked websites count
    blocked_websites_dict = admin_db.reference('blocked_sites').get() or {}
    blocked_websites_count = len(blocked_websites_dict)

    # Count connected devices from senior users
    connected_devices_count = sum(1 for user in senior_users_dict.values() if user.get('status') == 'online')

    # Fetch notification alerts count
    notifications_ref = admin_db.reference('notification_alerts')
    notifications = notifications_ref.get() or {}
    active_alerts_count = len(notifications)

    # Prepare recent activity
    recent_activity = []
    blocked_sites = [entry['website_url'] for entry in blocked_websites_dict.values()]

    for user_id, user_info in senior_users_dict.items():
        fullname = user_info.get('fullname', 'Unknown')
        history_sites = user_info.get('history_sites', {})

        for site_id, site_info in history_sites.items():
            website = site_info.get('url', '')
            timestamp = site_info.get('last_visited', '')
            status = 'Blocked' if website in blocked_sites else 'Safe'

            recent_activity.append({
                'fullname': fullname,
                'website': website,
                'time': timestamp,
                'status': status
            })

    try:
        recent_activity.sort(key=lambda x: datetime.strptime(x['time'], '%Y-%m-%d %H:%M:%S'), reverse=True)
    except:
        pass

    recent_activity = recent_activity[:5]

    return render_template('admin_dashboard.html',
                           senior_user_count=senior_user_count,
                           blocked_websites_count=blocked_websites_count,
                           connected_devices_count=connected_devices_count,
                           active_alerts_count=active_alerts_count,
                           recent_activity=recent_activity)


@app.route('/admin_senior_users')
def admin_senior_users():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403

    if 'user' not in session or session['user'].get('role') != 'admin':
        flash("Only Admin Can Access This Senior Users.", "error")
        return redirect(url_for('admin_dashboard'))

    senior_users_dict = admin_db.reference('senior_users').get() or {}

    senior_users = list(senior_users_dict.values())

    return render_template('admin_senior_users.html', senior_users=senior_users)


@app.route('/add_senior_user', methods=['POST'])
def add_senior_user():
    data = request.get_json()

    fullname = data.get('fullname', '').strip()
    email = data.get('email', '').strip()
    password = data.get('password', '').strip()
    mac_address = data.get('mac_address', '').strip().lower()

    if not all([fullname, email, password, mac_address]):
        return jsonify({'message': 'All fields are required.'}), 400

    mac_query = mac_address.replace(':', '-')
    if not mac_query or mac_query == '00-00-00-00-00-00':
        return jsonify({'message': 'Invalid MAC address.'}), 400

    # Fetch device name from MAC Vendors API
    try:
        r = requests.get(f'https://api.macvendors.com/{mac_query}')
        device_name = r.text if r.status_code == 200 else "Unknown Device"
    except Exception as e:
        device_name = "Unknown Device"
        print(f"[!] Error fetching device name: {e}")

    client_ip = request.remote_addr or "0.0.0.0"

    user_id = str(uuid.uuid4())

    user_data = {
        "fullname": fullname,
        "email": email,
        "password": password,
        "role": "senior",
        "status": "Offline",
        "filter_level": "Easy",
        "last_active": "Never",
        "device": device_name,  # Use the fetched device name
        "ip_address": client_ip,
        "mac_address": mac_address if mac_address else "MAC Not Found"
    }

    try:
        ref = admin_db.reference('senior_users')
        ref.child(user_id).set(user_data)
        return jsonify({"message": "User  added successfully"})
    except Exception as e:
        return jsonify({"message": "Failed to add user: " + str(e)}), 500


def log_browsing_activity():
    browsing_activity = fetch_browsing_activity()
    for activity in browsing_activity:
        # Assuming activity is a dictionary with user, website, and timestamp
        activity_ref = db.reference('website_activity')
        activity_ref.push(activity)


@app.route('/device_connected', methods=['POST'])
def device_connected():
    # This endpoint should be called when a device connects to the network
    log_browsing_activity()
    return jsonify({"message": "Browsing activity logged successfully"}), 200


@app.route('/log_activity', methods=['POST'])
def log_activity():
    data = request.get_json()
    user = data.get('user')
    website = data.get('website')
    ip_address = request.remote_addr
    timestamp = data.get('timestamp')

    # Validate the website URL
    if not website.startswith(('http://', 'https://')):
        website = 'http://' + website

    user_uid = session.get('user').get('uid')  # Get the UID from the session

    activity_ref = db.reference(f'website_activity/{user_uid}')
    activity_ref.push({
        'user': user,
        'website': website,
        'ip': ip_address,
        'timestamp': timestamp
    })
    return jsonify({"message": "Activity logged successfully"}), 200


@app.route('/admin_website_activity', methods=['GET'])
def admin_website_activity():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session:
        return redirect(url_for('login'))
    user_role = session['user'].get('role')
    if user_role not in ['admin', 'Guardian']:
        flash("Unauthorized access.", "error")
        return redirect(url_for('login'))

    # Fetch query parameters
    search = request.args.get('search', '').lower()
    status_filter = request.args.get('status', 'All Statuses')
    user_filter = request.args.get('user', 'All Users')

    # Fetch senior users from Firebase
    senior_users_ref = db.reference('senior_users')
    senior_users = senior_users_ref.get()

    # Fetch blocked sites from Firebase
    blocked_sites_ref = db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or []

    logs_list = []
    user_options = []

    for user_id, user_info in senior_users.items():
        fullname = user_info.get('fullname', 'Unknown')
        user_options.append(fullname)
        if user_filter != 'All Users' and fullname != user_filter:
            continue  # Skip non-matching users

        ip_address = user_info.get('ip_address', 'Unknown IP')
        history_sites = user_info.get('history_sites', {})

        for site_id, site_info in history_sites.items():
            website = site_info.get('url', '')
            action = 'visited'
            time = site_info.get('timestamp', '')
            status = 'Blocked' if website in blocked_sites else 'Allowed'

            # Apply filters
            if search and search not in website.lower() and search not in fullname.lower():
                continue
            if status_filter != 'All Statuses' and status_filter != status:
                continue

            logs_list.append({
                'fullname': fullname,
                'website': website,
                'action': action,
                'time': time,
                'status': status,
                'ip': ip_address
            })

    # Sort logs_list by time in descending order using dateutil.parser
    logs_list.sort(key=lambda x: parser.parse(x['time']), reverse=True)

    return render_template('admin_website_activity.html', logs=logs_list, users=user_options)


@app.route('/export_data', methods=['GET'])
def export_data():
    # Fetch logs from Firebase
    logs_ref = db.reference('website_activity')
    logs = logs_ref.get()

    # Create a new Document
    doc = Document()
    doc.add_heading('Senior User Website Activity', level=1)

    # Add table to the document
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Website'
    hdr_cells[2].text = 'Action'
    hdr_cells[3].text = 'Time'
    hdr_cells[4].text = 'Status'
    hdr_cells[5].text = 'IP Address'

    # Process logs and add to the table
    if logs:
        for key, value in logs.items():
            row_cells = table.add_row().cells
            row_cells[0].text = value.get('user', 'Unknown')
            row_cells[1].text = value.get('website', '')
            row_cells[2].text = 'visited'
            row_cells[3].text = value.get('timestamp', '')
            row_cells[4].text = 'Blocked' if value.get('website') in blocked_sites else 'Allowed'
            row_cells[5].text = value.get('ip', 'Unknown IP')

    # Save the document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Send the file to the user
    return send_file(doc_io, as_attachment=True, download_name='Senior_User_Website_Activity.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def fetch_browsing_activity():
    router_url = "https://tplinkwifi.net/webpages/index.html?t=e6bfcc9c"  # Use your router's URL
    session = requests.Session()
    session.verify = False  # Disable SSL verification for local connections

    # Replace with your router's login credentials
    username = "admin"
    password = "Admin12345"

    try:
        # Step 1: Log in to the router
        login_payload = {
            "username": username,
            "password": password
        }

        login_resp = session.post(f"{router_url}/login", data=login_payload)  # Adjusted endpoint

        if login_resp.status_code != 200 or "token" not in login_resp.json():
            print("[!] Login failed: No session token received.")
            return []

        # Step 2: Fetch browsing activity
        activity_resp = session.get(f"{router_url}/browsing_activity")  # Adjusted endpoint

        if activity_resp.status_code != 200:
            print("[!] Failed to fetch browsing activity.")
            return []

        # Assuming the response is in JSON format
        browsing_activity = activity_resp.json()

        # Process the browsing activity to extract relevant information
        activity_list = []
        for entry in browsing_activity.get("data", []):
            activity_list.append({
                'user': entry.get('user', 'Unknown'),
                'website': entry.get('url', ''),
                'timestamp': entry.get('timestamp', ''),
                'ip': entry.get('ip', '')
            })

        return activity_list

    except Exception as e:
        print(f"[!] Error fetching browsing activity: {e}")
        return []


@app.route('/admin_filtering_rules')
def admin_filtering_rules():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session or session['user'].get('role') != 'admin':
        flash("Only Admin Can Access This Filtering Rules.", "error")
        return redirect(url_for('admin_dashboard'))

    # Blocklist
    blocked_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_ref.get() or {}
    blocked_list = [{
        'id': key,
        'website_url': val.get('website_url'),
        'category': val.get('category'),
        'date_added': val.get('date_added')
    } for key, val in blocked_sites.items()]

    # Whitelisted sites
    whitelist_ref = admin_db.reference('whitelisted_sites')
    whitelisted_sites = whitelist_ref.get() or {}
    whitelisted_list = [{
        'id': key,
        'website_url': val.get('website_url'),
        'category': val.get('category'),
        'date_added': val.get('date_added')
    } for key, val in whitelisted_sites.items()]

    return render_template('admin_filtering_rules.html', blocked_list=blocked_list, whitelisted_list=whitelisted_list)


# Filtering Rules

@app.route('/add_blocked_site', methods=['POST'])
def add_blocked_site():
    if not check_admin_access():
        return redirect(url_for('login'))

    website_url = request.form['website_url']
    category = request.form['category']
    date_added = datetime.now(timezone.utc).strftime('%Y-%m-%d')

    new_entry = {
        'website_url': website_url.lower(),
        'category': category,
        'date_added': date_added
    }

    ref = admin_db.reference('blocked_sites')
    ref.push(new_entry)

    flash('Website blocked successfully!', 'success')
    return redirect(url_for('admin_filtering_rules'))


@app.route('/blocked_sites')
def get_blocked_sites():
    ref = admin_db.reference('blocked_sites')
    blocked = ref.get()
    if not blocked:
        return jsonify([])

    blocked_list = [entry['website_url'] for entry in blocked.values()]
    return jsonify(blocked_list)


@app.route('/delete_blocked_site/<site_id>', methods=['DELETE'])
def delete_blocked_site(site_id):
    if not check_admin_access():
        return "Unauthorized", 403

    ref = admin_db.reference(f'blocked_sites/{site_id}')
    ref.delete()
    return '', 204


@app.route('/add_whitelist', methods=['POST'])
def add_whitelist():
    if not check_admin_access():
        return redirect(url_for('login'))

    website_url = request.form.get('website_url')
    category = request.form.get('category')
    if not website_url or not category:
        flash("Missing data!", "error")
        return redirect(url_for('admin_filtering_rules'))

    entry_id = str(uuid.uuid4())
    whitelist_ref = admin_db.reference(f'whitelisted_sites/{entry_id}')
    whitelist_ref.set({
        'website_url': website_url,
        'category': category,
        'date_added': datetime.now(timezone.utc).strftime('%Y-%m-%d')
    })

    flash(f"Website '{website_url}' added to whitelist.", "success")
    return redirect(url_for('admin_filtering_rules'))


@app.route('/delete_whitelist/<site_id>', methods=['DELETE'])
def delete_whitelist(site_id):
    if not check_admin_access():
        return "Unauthorized", 403

    ref = admin_db.reference(f'whitelisted_sites/{site_id}')
    ref.delete()
    return '', 204


@app.route('/admin_devices')
def admin_devices():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session or session['user'].get('role') != 'admin':
        flash("Only Admin Can Access This Devices.", "error")
        return redirect(url_for('admin_dashboard'))

    # Fetch senior users data
    senior_users_dict = admin_db.reference('senior_users').get() or {}
    senior_users = list(senior_users_dict.values())

    # Count total, online, and offline devices
    total_devices = len(senior_users)
    online_devices = sum(1 for user in senior_users if user.get('status') == 'online')
    offline_devices = total_devices - online_devices

    return render_template('admin_devices.html', senior_users=senior_users,
                           total_devices=total_devices, online_devices=online_devices,
                           offline_devices=offline_devices)


@app.route('/admin_notification')
def admin_notification():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session:
        return redirect(url_for('login'))
    user_role = session['user'].get('role')
    if user_role not in ['admin', 'Guardian']:
        flash("Unauthorized access.", "error")
        return redirect(url_for('login'))

    # Fetch notification alerts from Firebase
    notifications_ref = admin_db.reference('notification_alerts')
    notifications = notifications_ref.get() or {}

    # Prepare the notifications for rendering
    notification_list = []
    for key, value in notifications.items():
        notification_list.append({
            'id': key,
            'date': value.get('date'),
            'name': value.get('name'),
            'time': value.get('time'),
            'type': value.get('type'),
            'type_risk': value.get('type_risk'),
            'website': value.get('website'),
            'email': value.get('name')  # Assuming 'name' is the email
        })

    # Fetch senior users for the filter
    senior_users_ref = admin_db.reference('senior_users')
    senior_users = senior_users_ref.get() or {}
    senior_users_list = [{'fullname': user.get('fullname'), 'email': user.get('email')} for user in senior_users.values()]

    return render_template('admin_notification.html', notifications=notification_list, senior_users=senior_users_list)


@app.route('/admin_report')
def admin_report():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session:
        return redirect(url_for('login'))
    user_role = session['user'].get('role')
    if user_role not in ['admin', 'Guardian']:
        flash("Unauthorized access.", "error")
        return redirect(url_for('login'))

    # Fetch blocked sites from Firebase
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}
    blocked_sites_count = len(blocked_sites)

    # Fetch senior users and their history sites
    senior_users_ref = admin_db.reference('senior_users')
    senior_users = senior_users_ref.get() or {}

    total_visits = 0
    most_active_seniors = []
    most_visited_sites_by_date = {}

    for user_id, user_info in senior_users.items():
        history_sites = user_info.get('history_sites', {})

        # Calculate total visits and most visited sites by date
        for site_id, site_info in history_sites.items():
            last_visited_date = site_info.get('last_visited', '').split(' ')[0]  # Get the date part
            visit_count = site_info.get('visit_count', 0)

            # Initialize the dictionary for the date if it doesn't exist
            if last_visited_date not in most_visited_sites_by_date:
                most_visited_sites_by_date[last_visited_date] = []

            # Append the site info to the corresponding date
            most_visited_sites_by_date[last_visited_date].append({
                'title': site_info.get('title', 'Unknown'),
                'url': site_info.get('url', ''),
                'visit_count': visit_count
            })

            total_visits += visit_count

        # Calculate the visit count for the user
        visit_count = sum(site_info.get('visit_count', 0) for site_info in history_sites.values())
        if visit_count > 0:
            most_active_seniors.append({
                'fullname': user_info.get('fullname', 'Unknown'),
                'visit_count': visit_count
            })

    # Determine the most visited site for each date
    for date, sites in most_visited_sites_by_date.items():
        # Sort sites by visit count in descending order
        sites.sort(key=lambda x: x['visit_count'], reverse=True)
        most_visited_sites_by_date[date] = sites[0] if sites else None  # Get the most visited site or None

    # Sort seniors by visit count in descending order
    most_active_seniors.sort(key=lambda x: x['visit_count'], reverse=True)
    most_active_senior = most_active_seniors[0] if most_active_seniors else {'fullname': 'None', 'visit_count': 0}

    # Count blocked sites by category
    category_counts = {}
    for site_info in blocked_sites.values():
        category = site_info.get('category', 'Uncategorized')
        category_counts[category] = category_counts.get(category, 0) + 1

    # Prepare data for most blocked categories
    most_blocked_categories = []
    for category, count in category_counts.items():
        most_blocked_categories.append({
            'category': category,
            'count': count
        })

    # Determine the top blocked category
    top_blocked_category = max(most_blocked_categories, key=lambda x: x['count'],
                               default={'category': 'None', 'count': 0})

    return render_template('admin_report.html',
                           blocked_sites_count=blocked_sites_count,
                           total_visits=total_visits,
                           most_visited_sites_by_date=most_visited_sites_by_date,
                           most_active_senior=most_active_senior,
                           most_active_seniors=most_active_seniors,  # <-- Add this line
                           top_blocked_category=top_blocked_category,
                           most_blocked_categories=most_blocked_categories)


@app.route('/export_reports', methods=['GET'])
def export_reports():
    # Create a new Document
    doc = Document()
    doc.add_heading('Exported Reports', level=1)

    # Fetch senior users from Firebase
    senior_users_ref = admin_db.reference('senior_users')
    senior_users = senior_users_ref.get() or {}

    # Fetch blocked sites from Firebase
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}

    # Fetch whitelisted sites from Firebase
    whitelisted_sites_ref = admin_db.reference('whitelisted_sites')
    whitelisted_sites = whitelisted_sites_ref.get() or {}

    # Add Senior Users Table
    doc.add_heading('Senior Users', level=2)
    senior_table = doc.add_table(rows=1, cols=6)
    senior_table.style = 'Table Grid'
    hdr_cells = senior_table.rows[0].cells
    hdr_cells[0].text = 'Full Name'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Status'
    hdr_cells[3].text = 'Connection Duration'
    hdr_cells[4].text = 'IP Address'
    hdr_cells[5].text = 'History Sites'

    for user_id, user_info in senior_users.items():
        row_cells = senior_table.add_row().cells
        row_cells[0].text = user_info.get('fullname', 'Unknown')
        row_cells[1].text = user_info.get('email', 'Unknown')
        row_cells[2].text = user_info.get('status', 'Unknown')
        row_cells[
            3].text = f"{user_info.get('connection_duration', {}).get('hours', 0)}h {user_info.get('connection_duration', {}).get('minutes', 0)}m {user_info.get('connection_duration', {}).get('seconds', 0)}s"
        row_cells[4].text = user_info.get('ip_address', 'Unknown')

        # Add history sites
        history_sites = user_info.get('history_sites', {})
        history_list = ', '.join(
            [f"{site['url']} (Visited: {site['last_visited']})" for site in history_sites.values()])
        row_cells[5].text = history_list

    # Add Blocked Sites Table
    doc.add_heading('Blocked Sites', level=2)
    blocked_table = doc.add_table(rows=1, cols=3)
    blocked_table.style = 'Table Grid'
    hdr_cells = blocked_table.rows[0].cells
    hdr_cells[0].text = 'Website URL'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Date Added'

    for site_info in blocked_sites.values():
        row_cells = blocked_table.add_row().cells
        row_cells[0].text = site_info.get('website_url', 'Unknown')
        row_cells[1].text = site_info.get('category', 'Unknown')
        row_cells[2].text = site_info.get('date_added', 'Unknown')

    # Add Whitelisted Sites Table
    doc.add_heading('Whitelisted Sites', level=2)
    whitelist_table = doc.add_table(rows=1, cols=3)
    whitelist_table.style = 'Table Grid'
    hdr_cells = whitelist_table.rows[0].cells
    hdr_cells[0].text = 'Website URL'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Date Added'

    for site_info in whitelisted_sites.values():
        row_cells = whitelist_table.add_row().cells
        row_cells[0].text = site_info.get('website_url', 'Unknown')
        row_cells[1].text = site_info.get('category', 'Unknown')
        row_cells[2].text = site_info.get('date_added', 'Unknown')

    # Save the document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Send the file to the user
    return send_file(doc_io, as_attachment=True, download_name='Exported_Reports.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/admin_setting', methods=['GET', 'POST'])
def admin_setting():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if 'user' not in session or session['user'].get('role') != 'admin':
        flash("Only Admin Can Access This Settings.", "error")
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        # Save notification preferences
        realtime_alerts = request.form.get('realtime', 'off') == 'on'
        email_notifications = request.form.get('emailnotif', 'off') == 'on'
        user_email = request.form.get('notification-email')

        # Save preferences to Firebase
        user_ref = admin_db.reference('admin_users/{}'.format(session['user']['id']))
        user_ref.update({
            'realtime_alerts': realtime_alerts,
            'email_notifications': email_notifications,
            'notification_email': user_email
        })

        flash("Notification settings updated successfully.", "success")
        return redirect(url_for('admin_setting'))

    # Fetch current admin users from Firebase
    admin_users_ref = admin_db.reference('admin_users')
    current_users = admin_users_ref.get() or {}

    # Convert to a list for easier rendering
    current_users_list = []
    for user_id, user_info in current_users.items():
        current_users_list.append({
            'email': user_info.get('email'),
            'role': user_info.get('role')
        })

    return render_template('admin_setting.html', active_tab='user-management', current_users=current_users_list)


def send_email_notification(to_email, subject, message):
    msg = MIMEText(message)
    msg['Subject'] = subject
    msg['From'] = 'your-email@example.com'
    msg['To'] = to_email

    with smtplib.SMTP('smtp.example.com', 587) as server:
        server.starttls()
        server.login('your-email@example.com', 'your-email-password')
        server.send_message(msg)


def notify_users(notification):
    # Fetch users from Firebase
    admin_users_ref = admin_db.reference('admin_users')
    current_users = admin_users_ref.get() or {}

    # Specify the email address to send notifications to
    target_email = "goldenshield04@gmail.com"

    for user_id, user_info in current_users.items():
        # Check if the user has email notifications enabled and if their email matches the target email
        if user_info.get('email_notifications') and user_info.get('notification_email') == target_email:
            send_email_notification(user_info['notification_email'], "New Notification", notification)

        if user_info.get('realtime_alerts'):
            # Code to send real-time alert to the admin dashboard
            pass  # Implement real-time alert logic here

# Add your notification creation logic here


@app.route('/create_notification', methods=['POST'])
def create_notification():
    # Get notification data from request
    notification_data = request.get_json()

    # Add timestamp for ordering
    notification_data['timestamp'] = firebase_admin.db.ServerValue.TIMESTAMP

    # Push to Firebase
    ref = admin_db.reference('notification_alerts')
    new_notification = ref.push(notification_data)

    return jsonify({"success": True, "id": new_notification.key})


@app.route('/update_filtering_preferences', methods=['POST'])
def update_filtering_preferences():
    data = request.get_json()
    block_gambling = data.get('block_gambling', False)
    block_adult = data.get('block_adult', False)

    # Update the preferences in the database (Firebase)
    admin_db.reference('filtering_preferences').update({
        'block_gambling': block_gambling,
        'block_adult': block_adult
    })

    return jsonify({"message": "Preferences updated successfully."}), 200


@app.route('/get_filtering_preferences', methods=['GET'])
def get_filtering_preferences():
    # Fetch filtering preferences from Firebase
    filtering_preferences_ref = admin_db.reference('filtering_preferences')
    filtering_preferences = filtering_preferences_ref.get() or {}

    return jsonify({
        'block_gambling': filtering_preferences.get('block_gambling', False),
        'block_adult': filtering_preferences.get('block_adult', False)
    })


@app.route('/delete_user/<email>', methods=['DELETE'])
def delete_user(email):
    # Fetch the user data from Firebase
    admin_users_ref = admin_db.reference('admin_users')
    users = admin_users_ref.get() or {}

    # Find the user by email and delete if not an admin
    for user_id, user_info in users.items():
        if user_info.get('email') == email:
            if user_info.get('role') == 'admin':
                return jsonify({"message": "Admin users cannot be deleted."}), 403
            else:
                admin_users_ref.child(user_id).delete()
                return jsonify({"message": "User  deleted successfully."}), 200

    return jsonify({"message": "User  not found."}), 404


@app.route('/update_password', methods=['POST'])
def update_password():
    email = request.form.get('email')
    new_password = request.form.get('new_password')

    # Fetch the user data from Firebase
    admin_users_ref = admin_db.reference('admin_users')
    users = admin_users_ref.get() or {}

    # Find the user by email and update the password
    for user_id, user_info in users.items():
        if user_info.get('email') == email:
            admin_users_ref.child(user_id).update({"password": new_password})
            flash("Password updated successfully!", "success")
            return redirect(url_for('admin_setting'))

    flash("User  not found.", "error")
    return redirect(url_for('admin_setting'))


@app.route('/add_user', methods=['POST'])
def add_user():
    email = request.form.get('email')
    role = request.form.get('role')
    password = "admin123"  # Default password

    # Generate a unique user ID
    user_id = str(uuid.uuid4())

    # User data to be added
    user_data = {
        "email": email,
        "fullname": email.split('@')[0],  # Use the part before @ as fullname
        "password": password,
        "role": role,
        "status": "online",
        "last_active": datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')
    }

    # Add user to the admin_users database
    admin_db.reference(f'admin_users/{user_id}').set(user_data)

    flash("User  added successfully!", "success")
    return redirect(url_for('admin_setting'))


def time_since_last_active(last_active_str):
    if last_active_str == "Never" or not last_active_str:
        return "Never"

    try:
        last_active_dt = datetime.fromisoformat(last_active_str.replace('Z', '+00:00'))
    except Exception:
        return "Unknown"

    now = datetime.now(timezone.utc)
    delta = now - last_active_dt

    seconds = delta.total_seconds()
    if seconds < 60:
        return f"{int(seconds)} seconds ago"
    elif seconds < 3600:
        minutes = int(seconds // 60)
        return f"{minutes} minutes ago"
    elif seconds < 86400:
        hours = int(seconds // 3600)
        return f"{hours} hours ago"
    else:
        days = int(seconds // 86400)
        return f"{days} days ago"


# ========== Senior User Routes ==========

def check_senior_access():
    if 'user' not in session or session['user'].get('role') != 'senior':
        flash("Unauthorized access.", "error")
        return False
    return True


@app.route('/log_browsing_history', methods=['POST'])
def log_browsing_history():
    data = request.get_json()
    user_email = data.get('user_email')  # Get the user's email from the request
    browsing_history = data.get('browsing_history', [])  # Get the browsing history

    if not user_email or not browsing_history:
        return jsonify({"message": "User  email and browsing history are required."}), 400

    # Save each browsing history entry to Firebase
    for entry in browsing_history:
        website = entry.get('url')
        timestamp = entry.get('timestamp', datetime.now(timezone.utc).isoformat())

        # Create a reference path based on user email
        activity_ref = db.reference(f'website_activity/{user_email}')

        # Push the activity to the user's specific path
        activity_ref.push({
            'website': website,
            'timestamp': timestamp
        })

    return jsonify({"message": "Browsing history logged successfully."}), 200


@app.route('/user_dashboard')
def user_dashboard():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return jsonify({"error": "Access Denied"}), 403
    if not check_senior_access():
        return jsonify({"error": "Unauthorized access"}), 403

    user_session = session.get('user')
    user_uid = user_session.get('uid')

    # Use get_history() to fetch browsing history
    browsing_history = get_history()

    # Log browsing history to Firebase
    log_browsing_history_to_firebase(user_uid, browsing_history)

    # Get initial history for display
    initial_history = browsing_history[:3]

    return render_template('user_dashboard.html', initial_history=initial_history, full_history=browsing_history)


@app.route('/check_access', methods=['POST'])
def check_access():
    data = request.get_json()
    website = data.get('website')

    if not website:
        return jsonify({"error": "Website URL is required."}), 400

    # Fetch blocked sites from Firebase
    blocked_sites_ref = admin_db.reference('blocked_sites')
    blocked_sites = blocked_sites_ref.get() or {}

    # Create a list of blocked site URLs
    blocked_list = [entry['website_url'].lower() for entry in blocked_sites.values()]

    # Fetch malicious sites from Firebase (assuming you have a separate reference for malicious sites)
    malicious_sites_ref = admin_db.reference('malicious_sites')
    malicious_sites = malicious_sites_ref.get() or {}

    # Create a list of malicious site URLs
    malicious_list = [entry['website_url'].lower() for entry in malicious_sites.values()]

    if not website:
        return jsonify({"error": "Website URL is required."}), 400
    # Check if the requested website is blocked
    if is_blocked_domain(website) or is_blocked_url(website):
        return jsonify({"message": "Access denied. This site is blocked."}), 403

    # Check if the requested website is in the blocked list
    if website.lower() in blocked_list:
        # Update the user's safety status to "Moderate"
        user_session = session.get('user')
        user_uid = user_session.get('uid')
        admin_db.reference(f'senior_users/{user_uid}').update({
            "filter_level": "Moderate"
        })
        return jsonify({"message": "Access denied. Safety status updated to Moderate."}), 403

    # Check if the requested website is in the malicious list
    if website.lower() in malicious_list:
        # Update the user's safety status to "Hard"
        user_session = session.get('user')
        user_uid = user_session.get('uid')
        admin_db.reference(f'senior_users/{user_uid}').update({
            "filter_level": "Hard"
        })
        return jsonify({"message": "Access denied. Safety status updated to Hard."}), 403

    return jsonify({"message": "Access granted."}), 200


def get_history():
    history_items = []
    history_sources = {
        'Chrome': os.path.expanduser(r'~\AppData\Local\Google\Chrome\User Data\Default\History'),
        'Brave': os.path.expanduser(r'~\AppData\Local\BraveSoftware\Brave-Browser\User Data\Default\History'),
        'Edge': os.path.expanduser(r'~\AppData\Local\Microsoft\Edge\User Data\Default\History'),
        'Android Chrome': os.path.expanduser(r'~/Android/data/com.android.chrome/databases/chrome'),
        'iOS Safari': '/var/mobile/Library/Safari/History',  # Placeholder for iOS Safari history
    }

    for browser, history_path in history_sources.items():
        if os.path.exists(history_path):
            temp_copy = f'temp_history_{browser.lower().replace(" ", "_")}'
            try:
                if 'Android' in browser:
                    # For Android, specify the database file
                    history_file = os.path.join(history_path, 'History')
                    if os.path.exists(history_file):
                        shutil.copy2(history_file, temp_copy)
                    else:
                        continue  # Skip if the history file does not exist
                elif 'iOS' in browser:
                    # For iOS, we are using a placeholder path
                    print(f"iOS Safari history path: {history_path}")  # Display the path
                    continue  # Skip actual reading for now
                else:
                    shutil.copy2(history_path, temp_copy)

                conn = sqlite3.connect(temp_copy)
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT urls.url, urls.title, urls.visit_count, urls.last_visit_time 
                    FROM urls 
                    ORDER BY last_visit_time DESC 
                    LIMIT 4
                """)
                results = cursor.fetchall()
                for url, title, visit_count, last_visit_time in results:
                    time_converted = datetime(1601, 1, 1) + timedelta(microseconds=last_visit_time)
                    history_entry = {
                        'url': url,
                        'title': title,
                        'visit_count': visit_count,
                        'last_visited': time_converted.strftime('%Y-%m-%d %H:%M:%S'),
                        'browser': browser
                    }
                    history_items.append(history_entry)
                    # Insert into Firebase
                    insert_history(history_entry)
            except Exception as e:
                history_items.append({
                    'url': 'N/A',
                    'title': f'Could not read {browser} history',
                    'visit_count': 0,
                    'last_visited': str(e),
                    'browser': browser
                })
            finally:
                conn.close()
                os.remove(temp_copy)

    return history_items


def insert_history(history_entry):
    # Create a unique key for each history entry
    ref = db.reference('history_sites')
    new_entry_ref = ref.push({
        'url': history_entry['url'],
        'title': history_entry['title'],
        'visit_count': history_entry['visit_count'],
        'last_visited': history_entry['last_visited'],
        'timestamp': datetime.now().isoformat(),  # Store the current timestamp
    })
    print(f"Inserted history entry with ID: {new_entry_ref.key}")


def log_browsing_history_to_firebase(user_uid, browsing_history):
    """Log the browsing history to Firebase under the user's history_sites."""
    history_ref = db.reference(f'senior_users/{user_uid}/history_sites')

    for entry in browsing_history:
        history_ref.push({
            'url': entry['url'],
            'title': entry['title'],
            'last_visited': entry['last_visited'],
            'visit_count': entry['visit_count'],
            'timestamp': datetime.now(timezone.utc).isoformat()  # Use the correct datetime
        })


@app.route('/api/user_browsing_history')
def get_user_browsing_history():
    client_ip = request.remote_addr

    if not is_allowed_ip(client_ip):
        return jsonify({"error": "Access Denied"}), 403

    user_session = session.get('user')
    if not user_session or not check_senior_access():
        return jsonify({"error": "Unauthorized access"}), 403

    user_email = user_session.get('email')
    if not user_email:
        return jsonify({"error": "Email not found in session"}), 400

    user_ref = admin_db.reference('senior_users')
    user_data = user_ref.order_by_child('email').equal_to(user_email).get()

    if not user_data:
        return jsonify({"error": "User not found"}), 404

    user_id = next(iter(user_data))
    logs_ref = db.reference(f'website_activity/{user_id}')
    logs = logs_ref.get() or {}

    browsing_history = []
    for key, value in logs.items():
        raw_timestamp = value.get('timestamp', '')
        try:
            formatted_time = datetime.fromisoformat(raw_timestamp).strftime('%Y-%m-%d %H:%M:%S')
        except Exception:
            formatted_time = 'Invalid timestamp'

        browsing_history.append({
            'website': value.get('website', 'Unknown'),
            'timestamp': raw_timestamp,
            'time': formatted_time
        })

    browsing_history.sort(key=lambda x: x.get('timestamp', ''), reverse=True)

    return jsonify({"browsing_history": browsing_history})




@app.route('/settings')
def user_setting():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if not check_senior_access():
        return redirect(url_for('login'))
    return render_template('user_setting.html')


@app.route('/help')
def user_help():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if not check_senior_access():
        return redirect(url_for('login'))
    return render_template('user_help.html')


@app.route('/faq')
def user_FAQ():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if not check_senior_access():
        return redirect(url_for('login'))
    return render_template('user_faq.html')


@app.route('/contact_us')
def user_contact_us():
    client_ip = request.remote_addr
    if not is_allowed_ip(client_ip):
        return "Access Denied: Connect to the TP-Link Wi-Fi network (192.168.0.x).", 403
    if not check_senior_access():
        return redirect(url_for('login'))
    return render_template('user_contact_us.html')


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)